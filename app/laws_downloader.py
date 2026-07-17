"""
Загрузчик законодательства РК по сфере ИС с adilet.zan.kz.
Скачивает все законы, подзаконные акты, кодексы и международные договоры
в формате Word (.docx) и сохраняет в папку /laws/.
Поддерживает обновление: перезаписывает файл если на сайте изменилась дата.
"""

import re
import json
import time
import logging
import hashlib
from pathlib import Path
from datetime import datetime

import requests
import urllib3
from bs4 import BeautifulSoup

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

from paths import LAWS_DIR
INDEX_FILE = LAWS_DIR / "_index.json"

ADILET_BASE = "https://adilet.zan.kz"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120",
    "Accept-Language": "ru-RU,ru;q=0.9",
}

# ─── Полный реестр законов ────────────────────────────────────────────────────
LAWS_REGISTRY = [

    # ── Основные законы по ИС ────────────────────────────────────────────────
    {
        "id": "Z990000456_",
        "title": "Закон РК о товарных знаках, знаках обслуживания и наименованиях мест происхождения товаров",
        "short": "Закон о ТЗ",
        "category": "ИС_основные",
        "enacted": "1999-07-26",
    },
    {
        "id": "Z960000006_",
        "title": "Закон РК об авторском праве и смежных правах",
        "short": "Закон об авторском праве",
        "category": "ИС_основные",
        "enacted": "1996-06-10",
    },
    {
        "id": "Z990000427_",
        "title": "Патентный закон Республики Казахстан",
        "short": "Патентный закон",
        "category": "ИС_основные",
        "enacted": "1999-07-16",
    },
    {
        "id": "Z990000422_",
        "title": "Закон РК об охране селекционных достижений",
        "short": "Закон о селекции",
        "category": "ИС_основные",
        "enacted": "1999-07-13",
    },

    # ── Кодексы (разделы по ИС) ──────────────────────────────────────────────
    {
        "id": "K990000409_",
        "title": "Гражданский кодекс РК (Особенная часть)",
        "short": "ГК РК (Особенная часть)",
        "category": "Кодексы",
        "enacted": "1999-07-01",
    },
    {
        "id": "K1400000226",
        "title": "Уголовный кодекс Республики Казахстан",
        "short": "УК РК",
        "category": "Кодексы",
        "enacted": "2014-07-03",
    },
    {
        "id": "K1400000235",
        "title": "Кодекс РК об административных правонарушениях",
        "short": "КоАП РК",
        "category": "Кодексы",
        "enacted": "2014-07-05",
    },
    {
        "id": "K1500000377",
        "title": "Гражданский процессуальный кодекс РК",
        "short": "ГПК РК",
        "category": "Кодексы",
        "enacted": "2015-10-31",
    },

    # ── Международные договоры ───────────────────────────────────────────────
    {
        "id": "Z000000054_",
        "title": "О присоединении Республики Казахстан к Конвенции об охране интересов производителей фонограмм (Женевская конвенция)",
        "short": "Женевская конвенция о фонограммах",
        "category": "Международные_договоры",
        "enacted": "2000-01-01",
    },
    {
        "id": "Z990000407_",
        "title": "Парижская конвенция по охране промышленной собственности",
        "short": "Парижская конвенция",
        "category": "Международные_договоры",
        "enacted": "1999-01-01",
    },
    {
        "id": "Z020000332_",
        "title": "Соглашение ТРИПС (о торговых аспектах прав интеллектуальной собственности)",
        "short": "ТРИПС",
        "category": "Международные_договоры",
        "enacted": "2020-01-01",
    },
    {
        "id": "Z990000410_",
        "title": "Договор о патентной кооперации (PCT)",
        "short": "Договор PCT",
        "category": "Международные_договоры",
        "enacted": "1999-01-01",
    },
]


def load_index() -> dict:
    if INDEX_FILE.exists():
        with open(INDEX_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_index(index: dict):
    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)


def get_adilet_doc(doc_id: str) -> tuple[str, str, str]:
    """
    Возвращает (текст_закона, дата_изменения, прямая_ссылка_на_word).
    Пробует найти Word-ссылку на странице; если нет — возвращает HTML-текст.
    """
    url = f"{ADILET_BASE}/rus/docs/{doc_id}"
    try:
        r = requests.get(url, headers=HEADERS, timeout=60, verify=False)
        r.raise_for_status()
        r.encoding = "utf-8"
        soup = BeautifulSoup(r.text, "html.parser")

        # Ищем дату последнего изменения
        changed_date = ""
        for tag in soup.find_all(["span", "div", "p"], string=re.compile(r"\d{2}\.\d{2}\.\d{4}")):
            m = re.search(r"\d{2}\.\d{2}\.\d{4}", tag.get_text())
            if m:
                changed_date = m.group(0)
                break

        # Ищем прямую ссылку на Word/RTF
        word_url = ""
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if any(x in href.lower() for x in [".doc", ".rtf", "word", "export"]):
                word_url = href if href.startswith("http") else ADILET_BASE + href
                break

        # Основной текст документа
        content_div = (
            soup.find("div", {"class": re.compile(r"doc.?content|document|text", re.I)})
            or soup.find("div", {"id": re.compile(r"doc|text|content", re.I)})
            or soup.find("body")
        )
        text = content_div.get_text("\n", strip=True) if content_div else r.text[:50000]

        return text, changed_date, word_url
    except Exception as e:
        logger.error(f"Ошибка загрузки {doc_id}: {e}")
        return "", "", ""


def download_word_file(word_url: str, dest_path: Path) -> bool:
    """Скачивает Word-файл по прямой ссылке."""
    try:
        r = requests.get(word_url, headers=HEADERS, timeout=60, stream=True, verify=False)
        r.raise_for_status()
        with open(dest_path, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        return True
    except Exception as e:
        logger.error(f"Ошибка скачивания Word {word_url}: {e}")
        return False


def text_to_docx(title: str, text: str, dest_path: Path, source_url: str = ""):
    """Конвертирует текст закона в .docx с форматированием."""
    if not HAS_DOCX:
        dest_path.with_suffix(".txt").write_text(text, encoding="utf-8")
        return

    doc = Document()

    # Заголовок
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Мета-информация
    meta = doc.add_paragraph()
    meta.add_run(f"Источник: {source_url}").italic = True
    meta.add_run(f"\nДата загрузки: {datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True
    doc.add_paragraph()

    # Текст по абзацам
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Статьи выделяем как заголовки
        if re.match(r"^(Статья|СТАТЬЯ|Глава|ГЛАВА|Раздел|РАЗДЕЛ)\s+\d+", line):
            p = doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

    doc.save(str(dest_path))


def download_law(law: dict, index: dict, force: bool = False) -> dict:
    """
    Скачивает один закон. Возвращает обновлённую запись индекса.
    """
    doc_id = law["id"]
    category_dir = LAWS_DIR / law["category"]
    category_dir.mkdir(parents=True, exist_ok=True)

    safe_name = re.sub(r'[^\w\s-]', '', law["short"]).strip().replace(" ", "_")
    docx_path = category_dir / f"{safe_name}.docx"
    source_url = f"{ADILET_BASE}/rus/docs/{doc_id}"

    print(f"  [{law['category']}] {law['short']}...", end=" ", flush=True)

    text, changed_date, word_url = get_adilet_doc(doc_id)

    if not text and not word_url:
        print("❌ не найден на Adilet")
        return index.get(doc_id, {"status": "not_found", "doc_id": doc_id})

    # Проверяем нужно ли обновление
    existing = index.get(doc_id, {})
    if not force and existing.get("changed_date") == changed_date and docx_path.exists():
        print("✓ актуален")
        return existing

    # Скачиваем Word если есть прямая ссылка
    saved_as_word = False
    if word_url:
        saved_as_word = download_word_file(word_url, docx_path)

    # Иначе конвертируем текст в DOCX
    if not saved_as_word and text:
        text_to_docx(law["title"], text, docx_path, source_url)
        saved_as_word = docx_path.exists()

    if saved_as_word:
        print(f"✅ ({docx_path.stat().st_size // 1024} KB)")
    else:
        print("⚠️ сохранён как текст")

    record = {
        "doc_id": doc_id,
        "title": law["title"],
        "short": law["short"],
        "category": law["category"],
        "source_url": source_url,
        "file_path": str(docx_path.relative_to(LAWS_DIR.parent)),
        "changed_date": changed_date,
        "downloaded_at": datetime.now().isoformat(),
        "status": "ok" if saved_as_word else "error",
    }
    return record


def download_all_laws(force: bool = False, progress_callback=None) -> dict:
    """
    Основная функция: скачивает все законы из реестра.
    Возвращает сводку {total, updated, skipped, errors}.
    """
    index = load_index()
    summary = {"total": len(LAWS_REGISTRY), "updated": 0, "skipped": 0, "errors": 0}

    print(f"\n{'='*60}")
    print(f"Загрузка законодательства РК по ИС")
    print(f"Всего документов: {len(LAWS_REGISTRY)}")
    print(f"{'='*60}\n")

    for i, law in enumerate(LAWS_REGISTRY):
        if progress_callback:
            progress_callback(i, len(LAWS_REGISTRY), law["short"])
        try:
            record = download_law(law, index, force=force)
            index[law["id"]] = record
            if record.get("status") == "ok":
                summary["updated"] += 1
            elif record.get("status") == "not_found":
                summary["errors"] += 1
            else:
                summary["skipped"] += 1
        except Exception as e:
            logger.error(f"Критическая ошибка {law['id']}: {e}")
            summary["errors"] += 1
        time.sleep(1)  # Уважаем сервер

    save_index(index)

    print(f"\n{'='*60}")
    print(f"Готово: обновлено {summary['updated']}, "
          f"пропущено {summary['skipped']}, ошибок {summary['errors']}")
    print(f"Файлы: {LAWS_DIR}")
    print(f"{'='*60}\n")

    return summary


def check_updates() -> list[dict]:
    """
    Проверяет обновления без скачивания.
    Возвращает список законов у которых изменилась дата.
    """
    index = load_index()
    updated = []
    for law in LAWS_REGISTRY:
        _, changed_date, _ = get_adilet_doc(law["id"])
        existing = index.get(law["id"], {})
        if changed_date and existing.get("changed_date") != changed_date:
            updated.append({**law, "old_date": existing.get("changed_date", "—"), "new_date": changed_date})
        time.sleep(0.5)
    return updated


if __name__ == "__main__":
    import sys
    force = "--force" in sys.argv
    download_all_laws(force=force)
