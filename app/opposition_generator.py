"""
Генератор документа «Возражение» в Апелляционный совет МЮ РК.
Шаблон основан на деле ТОО «Көркем Телеком» против ТОО «Сергек Медиа».
"""

from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Реквизиты клиента и представителя ──────────────────────────────────────

CLIENT = {
    "name":    "ТОО «Көркем Телеком»",
    "bin":     "110340003601",
    "address": "Казахстан, Алматинская область, город Қонаев, микрорайон 5, здание 3А",
    "iik":     "KZ88601A871000257451",
    "bank":    "АО «Народный банк Казахстана»",
    "bik":     "HSBKKZKX",
}

REPRESENTATIVE = {
    "name":    "Канзадаева Лаура Камиловна",
    "iin":     "000725600853",
    "address": "г. Астана, пр. Туран 9Б",
    "phone":   "+7 700 982 29 10",
    "email":   "l.kanzadayeva@sergekgroup.kz",
}

# Наши знаки (более ранние)
OUR_MARKS = [
    {
        "number":        "56289",
        "name":          "«СЕРГЕК»",
        "priority_date": "23.08.2016",
        "reg_date":      "15.06.2017",
        "expiry_date":   "23.08.2036",
        "classes":       [9],
    },
    {
        "number":        "62753",
        "name":          "«SERGEK»",
        "priority_date": "22.11.2017",
        "reg_date":      "10.01.2019",
        "expiry_date":   "22.11.2027",
        "classes":       [9],
    },
]

# ─── Вспомогательные функции форматирования ─────────────────────────────────

def _set_font(run, size_pt: int, bold: bool = False, italic: bool = False, color: str = None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _para(doc, text: str, bold=False, italic=False, size=12,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, color=None) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_font(run, size, bold=bold, italic=italic, color=color)
    return p


def _heading(doc, text: str, level=1) -> None:
    sizes = {1: 13, 2: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, sizes.get(level, 12), bold=True)


def _bullet(doc, text: str, italic=False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, 12, italic=italic)


# ─── Основная функция генерации ──────────────────────────────────────────────

def generate_opposition(
    contested_marks: list[dict],
    owner: dict | None = None,
    our_marks: list[dict] | None = None,
    client: dict | None = None,
    representative: dict | None = None,
    date_str: str | None = None,
) -> bytes:
    """
    Генерирует .docx возражения и возвращает байты файла.

    contested_marks: список оспариваемых знаков. Каждый словарь:
        {
            "name":          "Сергек AQTOBE",   # название
            "number":        "85439",            # номер свидетельства
            "app_number":    "113693",           # номер заявки
            "priority_date": "12.01.2023",       # дата приоритета
            "reg_date":      "15.08.2023",       # дата регистрации
            "expiry_date":   "12.01.2033",       # срок действия
            "classes":       [35, 38, 41],        # классы МКТУ
        }
    owner:   {"name": "ТОО «Сергек Медиа»", "address": "..."}
    our_marks: список наших знаков (по умолчанию OUR_MARKS)
    """
    cl  = client or CLIENT
    rep = representative or REPRESENTATIVE
    our = our_marks or OUR_MARKS
    own = owner or {"name": "[Владелец оспариваемых регистраций]", "address": ""}
    date_str = date_str or datetime.today().strftime("%d.%m.%Y")

    doc = Document()

    # Поля страницы A4
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    # Стиль по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Шапка (правый блок) ──────────────────────────────────────────────────
    def _right(text, bold=False, italic=False, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        _set_font(run, size, bold=bold, italic=italic)

    _right("В Апелляционный совет Министерства юстиции", bold=True)
    _right("Республики Казахстан", bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    _right("Заявитель:", bold=True)
    _right(cl["name"], bold=True)
    _right(f"БИН {cl['bin']}", italic=True)
    _right(f"адрес: {cl['address']}", italic=True)
    _right(f"ИИК {cl['iik']}", italic=True)
    _right(cl["bank"], italic=True)
    _right(f"БИК {cl['bik']}", italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    _right("Представитель Заявителя по доверенности:", bold=True)
    _right(rep["name"], italic=True)
    _right(f"ИИН {rep['iin']}", italic=True)
    _right(f"Адрес: {rep['address']}", italic=True)
    _right(f"Тел: {rep['phone']}", italic=True)
    _right(f"эл. адрес: {rep['email']}", italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    _right("Владелец оспариваемых регистраций:", bold=True)
    _right(own["name"], bold=True)
    if own.get("address"):
        _right(f"адрес: {own['address']}", italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Заголовок ────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after  = Pt(4)
    run = p_title.add_run("ВОЗРАЖЕНИЕ")
    _set_font(run, 14, bold=True)

    # Подзаголовок: против каких знаков
    marks_list_str = "; ".join(
        f"«{m['name']}» по свидетельству № {m['number']}" for m in contested_marks
    )
    owner_name = own["name"]
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    run = p_sub.add_run(
        f"против регистрации товарных знаков {marks_list_str}, "
        f"зарегистрированных на имя {owner_name}"
    )
    _set_font(run, 12, italic=True)

    # ── Вводный абзац ────────────────────────────────────────────────────────
    our_refs = " и ".join(
        f"№{m['number']} {m['name']}" for m in our
    )
    _para(doc,
        f"{cl['name']} (далее — «Заявитель») как правообладатель более ранних "
        f"товарных знаков {our_refs} заявляет возражение против регистрации "
        f"следующих товарных знаков, зарегистрированных на имя {owner_name} "
        f"(далее — «{owner_name}»):"
    )

    # Нумерованный список оспариваемых знаков
    for i, m in enumerate(contested_marks, 1):
        classes_str = ", ".join(str(c) for c in m.get("classes", []))
        text = (
            f"{i}) комбинированный товарный знак «{m['name']}» по свидетельству "
            f"№ {m['number']} (заявка № {m.get('app_number','—')}, "
            f"приоритет от {m['priority_date']}, "
            f"дата регистрации {m['reg_date']}, "
            f"срок действия регистрации до {m['expiry_date']}) "
            f"в отношении услуг {classes_str} классов МКТУ;"
        )
        _para(doc, text)

    _para(doc,
        "Регистрация оспариваемых товарных знаков подлежит признанию "
        "недействительной по следующим основаниям:", bold=True
    )

    grounds = [
        f"Заявитель обладает более ранними исключительными правами на "
        f"товарные знаки {' / '.join(m['name'] for m in our)};",
        f"Обозначение {' / '.join(m['name'] for m in our)} обладает высокой "
        f"узнаваемостью и устойчиво ассоциируется, а использование тождественных "
        f"знаков создаёт риск введения потребителей в заблуждение относительно "
        f"лица, оказывающего услуги.",
        "Отсутствие самостоятельной различительной способности оспариваемых обозначений.",
        "Заявитель не предоставлял права на использование интеллектуальной собственности.",
    ]
    for g in grounds:
        _bullet(doc, g, italic=True)

    # ── Раздел 1 ─────────────────────────────────────────────────────────────
    marks_slash = " / ".join(m["name"] for m in our)
    _heading(doc, f"Раздел 1. Заявитель обладает более ранними исключительными правами "
                  f"на товарные знаки {marks_slash}.")

    _para(doc,
        f"Заявитель является лицом, на системной и постоянной основе использующим, "
        f"эксплуатирующим и продвигающим на территории Республики Казахстан обозначение "
        f"{marks_slash}, под которым индивидуализируется аппаратно-программный комплекс "
        f"фото-видеофиксации нарушений правил дорожного движения и сопутствующая система "
        f"обеспечения дорожной безопасности (далее — «система «Сергек»)."
    )

    _para(doc,
        "Заявитель является правообладателем следующих более ранних товарных знаков, "
        "словесный элемент которых тождественен доминирующему элементу оспариваемых "
        "обозначений:"
    )

    # Таблица наших знаков
    tbl = doc.add_table(rows=1 + len(our), cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Знак"
    hdr[1].text = "Регистрация и приоритет"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    for i, m in enumerate(our):
        row = tbl.rows[i + 1].cells
        row[0].text = m["name"]
        row[1].text = (
            f"№{m['number']};\n"
            f"дата подачи/приоритет — {m['priority_date']};\n"
            f"дата регистрации — {m['reg_date']};\n"
            f"срок действия — до {m['expiry_date']}."
        )
        for cell in row:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _para(doc, "Правовое основание на товарные знаки", bold=True, italic=True)
    _para(doc,
        "Исключительные права на указанные товарные знаки, охраняемые на территории "
        "Республики Казахстан, были уступлены Заявителю на основании Договора уступки "
        "исключительных прав на товарные знаки от 18 июля 2025 года, зарегистрированного "
        "в РГП на ПХВ «Национальный институт интеллектуальной собственности» Комитета по "
        "правам интеллектуальной собственности Министерства юстиции Республики Казахстан "
        "под номером № 04-20251472/14-21 от 21 августа 2025 года (приложение)."
    )

    _para(doc, "Приоритет", bold=True, italic=True)
    priority_text_parts = []
    for m in our:
        priority_text_parts.append(
            f"приоритет товарного знака {m['name']} — {m['priority_date']}"
        )
    for cm in contested_marks:
        priority_text_parts.append(
            f"приоритет оспариваемого знака «{cm['name']}» — {cm['priority_date']}"
        )
    _para(doc,
        "Приоритеты товарных знаков Заявителя (" +
        "; ".join(m['priority_date'] for m in our) + ") существенно предшествуют "
        "приоритетам оспариваемых товарных знаков (" +
        "; ".join(cm['priority_date'] for cm in contested_marks) + "), что "
        "подтверждает наличие у Заявителя более ранних прав."
    )

    _para(doc, f"Использование обозначения {marks_slash} в гражданском обороте.",
          bold=True, italic=True)
    _para(doc,
        f"Заявитель использует обозначение {marks_slash} в гражданском обороте "
        "Республики Казахстан задолго до даты подачи заявок оспариваемых знаков. "
        "Фактическое использование подтверждается заключением договоров государственно-частного "
        "партнёрства, государственных закупок, публикациями в СМИ, упоминаниями в "
        "интернет-источниках и социальных сетях, маркетинговым исследованием."
    )

    _para(doc,
        "В силу статьи 23 Закона РК «О товарных знаках, знаках обслуживания, "
        "географических указаниях и наименованиях мест происхождения товаров» "
        "регистрация товарного знака может быть оспорена и признана недействительной "
        "полностью или частично в течение всего срока действия, если она осуществлена "
        "в нарушение требований статей 6 и 7 Закона."
    )

    # ── Раздел 2 ─────────────────────────────────────────────────────────────
    _heading(doc, f"Раздел 2. Обозначения {marks_slash} обладают высокой узнаваемостью.")

    _para(doc,
        f"Обозначение {marks_slash} к моменту подачи заявок на регистрацию оспариваемых "
        "товарных знаков уже обладало высокой степенью узнаваемости на территории "
        "Республики Казахстан и имело высокую ассоциацию с «Системой «Сергек» — основным "
        "направлением деятельности Товарищества, что подтверждается отчётом по результатам "
        "маркетингового исследования."
    )
    _para(doc,
        "Согласно отчёту по результатам маркетингового исследования, обозначение «SERGEK» "
        "известно 79,3% опрошенных респондентов. При этом большинство респондентов "
        "ассоциирует данное обозначение с деятельностью в сфере дорожной безопасности, "
        "видеонаблюдения и видеофиксации.", bold=True
    )
    _para(doc,
        f"Следовательно, использование третьим лицом в составе товарного знака тождественного "
        f"словесного элемента {marks_slash} объективно создаёт у потребителей неверную "
        "ассоциацию системы «Сергек» с владельцем обжалуемых товарных знаков."
    )

    # ── Раздел 3 ─────────────────────────────────────────────────────────────
    _heading(doc, "Раздел 3. Отсутствие самостоятельной различительной способности оспариваемых обозначений.")

    _para(doc,
        "Согласно подпункту 1) пункта 3 статьи 6 Закона не допускается регистрация "
        "обозначений, являющихся ложными или способными ввести в заблуждение относительно "
        "товара или его изготовителя, услуги или лица, предоставляющего услуги."
    )
    _para(doc,
        "Согласно сведениям из Государственного реестра товарных знаков РК, при регистрации "
        "оспариваемых товарных знаков экспертная организация признала географические "
        "элементы неохраноспособными на основании подпункта 6) пункта 1 статьи 6 Закона. "
        "Следовательно, единственным охраноспособным элементом оспариваемых знаков является "
        f"словесный элемент {marks_slash}, который полностью воспроизводит более ранние "
        f"товарные знаки Заявителя по свидетельствам № "
        + " и № ".join(m["number"] for m in our) + "."
    )
    _para(doc,
        "Таким образом, добавление к охраняемому обозначению неохраноспособного "
        "географического элемента не устраняет сходство оспариваемых товарных знаков "
        "с более ранними товарными знаками Заявителя до степени смешения."
    )

    # ── Раздел 4 ─────────────────────────────────────────────────────────────
    _heading(doc, f"Раздел 4. Отсутствие разрешения на использование обозначения {marks_slash}.")

    _para(doc,
        "В соответствии со статьями 1025, 1026 и 1030 Гражданского кодекса РК "
        "правообладателю принадлежит исключительное право пользования и распоряжения "
        "товарным знаком. Использование товарного знака третьими лицами допускается только "
        "при наличии разрешения правообладателя по лицензионному договору."
    )
    _para(doc,
        f"Между Заявителем и {owner_name} отсутствуют лицензионные, договорные, "
        "корпоративные либо иные правоотношения, предоставляющие последнему право "
        f"использовать обозначения {marks_slash} либо регистрировать товарные знаки, "
        "содержащие указанные словесные элементы. Заявитель согласия не предоставлял."
    )

    # ── Просительная часть ───────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _para(doc,
        "На основании изложенного, руководствуясь подпунктом 1) пункта 3 статьи 6, "
        "подпунктом 1) пункта 1 статьи 7, статьями 23 и 41-2 Закона Республики Казахстан "
        "«О товарных знаках, знаках обслуживания, географических указаниях и наименованиях "
        "мест происхождения товаров», ПРОСИМ:", bold=True
    )

    for m in contested_marks:
        classes_str = ", ".join(str(c) for c in m.get("classes", []))
        _bullet(doc,
            f"Признать недействительной полностью регистрацию товарного знака "
            f"№{m['number']} «{m['name']}», зарегистрированного на имя {owner_name}, "
            f"в отношении всех услуг {classes_str} классов МКТУ."
        )
    _bullet(doc,
        "Обязать экспертную организацию принять необходимые меры по исполнению "
        "принятого решения."
    )

    # ── Приложения ───────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _para(doc, "Приложения:", bold=True, italic=True)

    attachments = []
    for i, m in enumerate(contested_marks, 1):
        attachments.append(
            f"{i}) выписка из Государственного реестра товарных знаков РК "
            f"в отношении товарного знака «{m['name']}» по свидетельству № {m['number']};"
        )
    n = len(contested_marks) + 1
    for m in our:
        attachments.append(
            f"{n}) выписка из Государственного реестра товарных знаков РК "
            f"в отношении товарного знака {m['name']} по свидетельству № {m['number']};"
        )
        n += 1
    attachments += [
        f"{n}) Договор уступки исключительных прав на товарные знаки от 18.07.2025 "
        f"(регистрационный № 04-20251472/14-21 от 21.08.2025 года);",
        f"{n+1}) уведомление о регистрации передачи исключительного права;",
        f"{n+2}) отчёт по результатам маркетингового исследования;",
        f"{n+3}) доверенность представителя Заявителя.",
    ]
    for att in attachments:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(att)
        _set_font(run, 12, italic=True)

    # ── Подпись ──────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_date.add_run(f"«___» __________ {datetime.today().year} года")
    _set_font(run, 12)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_sign.add_run(f"Представитель Заявителя по доверенности: {rep['name']}")
    _set_font(run, 12)

    p_sig_line = doc.add_paragraph()
    p_sig_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_sig_line.add_run("_______________________  /________________/")
    _set_font(run, 12)

    # ── Сохраняем в байты ────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def opposition_filename(contested_marks: list[dict]) -> str:
    """Имя файла для скачивания."""
    nums = "_".join(m["number"] for m in contested_marks if m.get("number"))
    date = datetime.today().strftime("%Y-%m-%d")
    return f"Возражение_{nums}_{date}.docx"


# ─── AI-генерация ────────────────────────────────────────────────────────────

def _build_header(doc: Document, case_data: dict) -> None:
    """Формирует шапку документа (правый блок адресатов)."""
    def _right(text, bold=False, italic=False, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        _set_font(run, size, bold=bold, italic=italic)

    _right("В Апелляционный совет Министерства юстиции", bold=True)
    _right("Республики Казахстан", bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    applicant = case_data.get("applicant_name", "")
    if applicant:
        _right("Заявитель:", bold=True)
        _right(applicant, bold=True)
        if case_data.get("applicant_bin"):
            _right(f"БИН {case_data['applicant_bin']}", italic=True)
        if case_data.get("applicant_address"):
            _right(f"адрес: {case_data['applicant_address']}", italic=True)
        if case_data.get("applicant_iik"):
            _right(f"ИИК {case_data['applicant_iik']}", italic=True)
        if case_data.get("applicant_bank"):
            _right(case_data["applicant_bank"], italic=True)
        if case_data.get("applicant_bik"):
            _right(f"БИК {case_data['applicant_bik']}", italic=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    rep = case_data.get("rep_name", "")
    if rep:
        _right("Представитель Заявителя по доверенности:", bold=True)
        _right(rep, italic=True)
        if case_data.get("rep_iin"):
            _right(f"ИИН {case_data['rep_iin']}", italic=True)
        if case_data.get("rep_address"):
            _right(f"Адрес: {case_data['rep_address']}", italic=True)
        if case_data.get("rep_phone"):
            _right(f"Тел: {case_data['rep_phone']}", italic=True)
        if case_data.get("rep_email"):
            _right(f"эл. адрес: {case_data['rep_email']}", italic=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    owner = case_data.get("owner_name", "")
    if owner:
        _right("Владелец оспариваемых регистраций:", bold=True)
        _right(owner, bold=True)
        if case_data.get("owner_address"):
            _right(f"адрес: {case_data['owner_address']}", italic=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _build_title(doc: Document, contested_marks: list[dict], owner_name: str) -> None:
    """Заголовок «ВОЗРАЖЕНИЕ»."""
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run = p_title.add_run("ВОЗРАЖЕНИЕ")
    _set_font(run, 14, bold=True)

    marks_str = "; ".join(
        f"«{m['name']}» №{m['number']}" for m in contested_marks if m.get("name")
    )
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    run = p_sub.add_run(
        f"против регистрации товарных знаков {marks_str}"
        + (f", зарегистрированных на имя {owner_name}" if owner_name else "")
    )
    _set_font(run, 12, italic=True)


def _render_ai_body(doc: Document, ai_text: str) -> None:
    """
    Разбирает AI-текст с разделами ## ... и вставляет в документ.
    Строки, начинающиеся с ## — заголовки разделов.
    Строки, начинающиеся с - или • — маркированные пункты.
    Остальное — обычные абзацы.
    """
    import re

    for raw_line in ai_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        if line.startswith("## "):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text)
            _set_font(run, 12, bold=True)

        elif line.startswith(("- ", "• ", "– ", "— ")):
            body = re.sub(r'^[-•–—]\s+', '', line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run("— " + body)
            _set_font(run, 12)

        elif re.match(r'^\d+\)', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line)
            _set_font(run, 12)

        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(line)
            _set_font(run, 12)


def _build_signature(doc: Document, case_data: dict) -> None:
    """Блок подписи."""
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_date.add_run(f"«___» __________ {datetime.today().year} года")
    _set_font(run, 12)

    signer = case_data.get("rep_name") or case_data.get("applicant_name") or ""
    if signer:
        p_sign = doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label = "Представитель Заявителя по доверенности: " if case_data.get("rep_name") else "Заявитель: "
        run = p_sign.add_run(label + signer)
        _set_font(run, 12)

    p_sig_line = doc.add_paragraph()
    p_sig_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_sig_line.add_run("_______________________  /________________/")
    _set_font(run, 12)


def generate_opposition_from_ai_text(ai_text: str, case_data: dict) -> bytes:
    """
    Создаёт .docx из AI-сгенерированного текста возражения.

    ai_text: текст, возвращённый generate_opposition_text() из ai_analyzer.py
    case_data: тот же словарь, что передавался в generate_opposition_text()
    """
    doc = Document()

    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    _build_header(doc, case_data)
    _build_title(
        doc,
        case_data.get("contested_marks", []),
        case_data.get("owner_name", ""),
    )
    _render_ai_body(doc, ai_text)
    _build_signature(doc, case_data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Извлекает текст из .docx-файла."""
    try:
        doc = Document(BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Извлекает текст из .pdf-файла."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
        return "\n".join(parts)
    except Exception as e:
        return f"[PDF: не удалось извлечь текст — {e}]"


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Извлекает текст из .xlsx-файла (все листы, все ячейки)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"[Лист: {sheet.title}]\n" + "\n".join(rows))
        return "\n\n".join(parts)
    except Exception as e:
        return f"[XLSX: не удалось извлечь текст — {e}]"


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Универсальный экстрактор по расширению файла."""
    name = filename.lower()
    if name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if name.endswith((".xlsx", ".xls")):
        return extract_text_from_xlsx(file_bytes)
    return ""
