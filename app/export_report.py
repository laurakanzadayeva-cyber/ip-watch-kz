"""
Экспорт отчётов в CSV, Word и PDF.
"""
import csv
import io
from datetime import datetime
from pathlib import Path

RISK_LABELS = {
    "high": "Высокий",
    "medium": "Средний",
    "low": "Низкий",
    "informational": "Информационный",
}
LEGAL_LABELS = {
    "not_reviewed": "Не проверено",
    "risk_confirmed": "Риск подтверждён",
    "risk_not_confirmed": "Риск не подтверждён",
    "archived": "Архив",
}
SOURCE_LABELS = {
    "kz_registry": "Реестр KZ",
    "kz_bulletin": "Бюллетень KZ",
    "wipo": "WIPO",
    "madrid": "Madrid",
    "manual": "Добавлено вручную",
}
OBJECT_LABELS = {
    "trademark": "Товарный знак",
    "well_known": "Общеизвестный ТЗ",
    "trade_name": "Фирменное наименование",
}
STATUS_LABELS = {
    "active": "Действует",
    "application": "Заявка",
    "expired": "Прекращён",
    "refused": "Отказ",
    "unknown": "Иной",
}

_CSV_HEADERS = [
    "Риск", "Обозначение", "Тип объекта", "Источник",
    "№ заявки", "№ регистрации", "Дата заявки", "Дата регистрации",
    "Правообладатель", "Адрес правообладателя", "Классы МКТУ",
    "Статус знака", "Причина совпадения", "Юридический статус",
    "Комментарий юриста", "Рекомендация", "В отчёт", "Ссылка",
]


def _mark_to_row(mark: dict) -> list:
    classes_raw = mark.get("nice_classes_str", "") or mark.get("nice_classes", "")
    risk = mark.get("risk_level", "")
    return [
        RISK_LABELS.get(risk, risk),
        mark.get("designation", ""),
        OBJECT_LABELS.get(mark.get("object_type"), mark.get("object_type", "")),
        SOURCE_LABELS.get(mark.get("source_code"), mark.get("source_code", "")),
        mark.get("application_number", ""),
        mark.get("registration_number", ""),
        mark.get("application_date", ""),
        mark.get("registration_date", ""),
        mark.get("owner", ""),
        mark.get("owner_address", ""),
        str(classes_raw),
        STATUS_LABELS.get(mark.get("status_mark"), mark.get("status_mark", "")),
        mark.get("match_reason", ""),
        LEGAL_LABELS.get(mark.get("legal_status"), mark.get("legal_status", "")),
        mark.get("lawyer_comment", ""),
        mark.get("recommended_action", ""),
        "Да" if mark.get("include_in_report") else "Нет",
        mark.get("source_url", ""),
    ]


# ─── CSV ─────────────────────────────────────────────────────────────────────

def generate_csv(marks: list[dict], **_) -> bytes:
    """CSV с BOM — открывается в Excel без настройки кодировки."""
    buf = io.StringIO()
    writer = csv.writer(buf, dialect="excel")
    writer.writerow(_CSV_HEADERS)
    for m in marks:
        writer.writerow(_mark_to_row(m))
    return buf.getvalue().encode("utf-8-sig")


# ─── WORD ─────────────────────────────────────────────────────────────────────

def generate_word(
    marks: list[dict],
    title: str,
    period_from: str,
    period_to: str,
    profiles: list[str],
    sources: list[str],
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _cell_width(cell, cm_val: float):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(cm_val * 567)))  # 567 twips/cm
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ── Заголовок ──
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for text in [
        f"Период: {period_from} — {period_to}",
        f"Профили: {', '.join(profiles) or 'все'}",
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)

    doc.add_paragraph()

    # ── Сводка ──
    doc.add_heading("Сводка", level=1)

    stats_data = [
        ("Всего найдено", len(marks)),
        ("Высокий риск", sum(1 for m in marks if m.get("risk_level") == "high")),
        ("Средний риск", sum(1 for m in marks if m.get("risk_level") == "medium")),
        ("Низкий риск", sum(1 for m in marks if m.get("risk_level") == "low")),
        ("Информационные", sum(1 for m in marks if m.get("risk_level") == "informational")),
        ("Не проверено юристом", sum(1 for m in marks if m.get("legal_status") == "not_reviewed")),
        ("Риск подтверждён", sum(1 for m in marks if m.get("legal_status") == "risk_confirmed")),
        ("Включено в отчёт", sum(1 for m in marks if m.get("include_in_report"))),
    ]

    stbl = doc.add_table(rows=len(stats_data) + 1, cols=2)
    stbl.style = "Table Grid"
    stbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Заголовок сводки
    for ci, text in enumerate(["Показатель", "Значение"]):
        c = stbl.rows[0].cells[ci]
        c.text = text
        _cell_bg(c, "37474F")
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (label, value) in enumerate(stats_data, 1):
        stbl.rows[i].cells[0].text = label
        stbl.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # ── Таблица результатов ──
    doc.add_heading(f"Результаты мониторинга ({len(marks)} записей)", level=1)

    RISK_FILL_HEX = {
        "high": "FFEBEE",
        "medium": "FFF3E0",
        "low": "E8F5E9",
        "informational": "E3F2FD",
    }

    col_headers = ["Риск", "Обозначение", "Тип", "Источник", "№ рег.", "Дата рег.", "Правообладатель", "Юр. статус"]
    col_widths_cm = [2.0, 4.5, 2.5, 2.2, 2.5, 2.5, 4.5, 2.8]

    if marks:
        tbl2 = doc.add_table(rows=len(marks) + 1, cols=len(col_headers))
        tbl2.style = "Table Grid"

        # Заголовок таблицы
        for ci, hdr_text in enumerate(col_headers):
            c = tbl2.rows[0].cells[ci]
            c.text = hdr_text
            _cell_bg(c, "37474F")
            _cell_width(c, col_widths_cm[ci])
            run = c.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Строки данных
        for ri, mark in enumerate(marks, 1):
            risk = mark.get("risk_level", "informational")
            fill = RISK_FILL_HEX.get(risk, "FFFFFF")
            short_row = [
                RISK_LABELS.get(risk, risk),
                mark.get("designation", ""),
                OBJECT_LABELS.get(mark.get("object_type"), mark.get("object_type", "")),
                SOURCE_LABELS.get(mark.get("source_code"), mark.get("source_code", "")),
                mark.get("registration_number", ""),
                mark.get("registration_date", ""),
                (mark.get("owner") or "")[:60],
                LEGAL_LABELS.get(mark.get("legal_status"), mark.get("legal_status", "")),
            ]
            for ci, val in enumerate(short_row):
                c = tbl2.rows[ri].cells[ci]
                c.text = str(val)
                _cell_bg(c, fill)
                _cell_width(c, col_widths_cm[ci])
                run = c.paragraphs[0].runs[0]
                run.font.size = Pt(8)
    else:
        doc.add_paragraph("Нет записей по выбранным параметрам.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── PDF ─────────────────────────────────────────────────────────────────────

def _find_cyrillic_font() -> str | None:
    # Bundled в fpdf2 >= 2.7.4
    try:
        import fpdf as _fpdf_mod
        bundled = Path(_fpdf_mod.__file__).parent / "fonts" / "DejaVuSans.ttf"
        if bundled.exists():
            return str(bundled)
    except Exception:
        pass
    # Системные шрифты
    for path in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]:
        if Path(path).exists():
            return path
    return None


def generate_pdf(
    marks: list[dict],
    title: str,
    period_from: str,
    period_to: str,
    profiles: list[str],
    sources: list[str],
) -> bytes:
    from fpdf import FPDF

    font_path = _find_cyrillic_font()

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    if font_path:
        pdf.add_font("Main", fname=font_path)
        pdf.add_font("Main", style="B", fname=font_path)
        fn = "Main"
    else:
        fn = "Helvetica"

    # ── Заголовок ──
    pdf.set_font(fn, style="B", size=16)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.ln(3)

    pdf.set_font(fn, size=9)
    for line in [
        f"Период: {period_from} — {period_to}",
        f"Профили: {', '.join(profiles) or 'все'}",
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
    ]:
        pdf.cell(0, 6, line, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # ── Сводка ──
    pdf.set_font(fn, style="B", size=12)
    pdf.cell(0, 8, "Сводка", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    high = sum(1 for m in marks if m.get("risk_level") == "high")
    medium = sum(1 for m in marks if m.get("risk_level") == "medium")
    low = sum(1 for m in marks if m.get("risk_level") == "low")
    info = sum(1 for m in marks if m.get("risk_level") == "informational")
    not_rev = sum(1 for m in marks if m.get("legal_status") == "not_reviewed")

    pdf.set_font(fn, size=9)
    for label, value, color in [
        ("Всего найдено", len(marks), (33, 33, 33)),
        ("Высокий риск", high, (211, 47, 47) if high else (33, 33, 33)),
        ("Средний риск", medium, (230, 81, 0) if medium else (33, 33, 33)),
        ("Низкий риск", low, (56, 142, 60) if low else (33, 33, 33)),
        ("Информационные", info, (25, 118, 210) if info else (33, 33, 33)),
        ("Не проверено юристом", not_rev, (33, 33, 33)),
    ]:
        pdf.set_fill_color(245, 245, 245)
        pdf.cell(90, 7, label, border="LTB", fill=True, new_x="RIGHT", new_y="LAST")
        pdf.set_text_color(*color)
        pdf.cell(25, 7, str(value), border="RTB", fill=False, align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
    pdf.ln(8)

    # ── Таблица результатов ──
    pdf.set_font(fn, style="B", size=12)
    pdf.cell(0, 8, f"Результаты мониторинга ({len(marks)} записей)", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    if not marks:
        pdf.set_font(fn, size=10)
        pdf.cell(0, 8, "Нет записей.", new_x="LMARGIN", new_y="NEXT")
    else:
        W = pdf.w - pdf.l_margin - pdf.r_margin
        raw_ws = [22, 48, 24, 22, 28, 26]
        scale = W / sum(raw_ws)
        col_ws = [round(w * scale, 1) for w in raw_ws]
        col_headers = ["Риск", "Обозначение", "Источник", "№ рег.", "Дата рег.", "Правообладатель"]

        def _table_header():
            pdf.set_fill_color(55, 71, 79)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font(fn, style="B", size=8)
            for i, (hdr, w) in enumerate(zip(col_headers, col_ws)):
                nx = "RIGHT" if i < len(col_headers) - 1 else "LMARGIN"
                ny = "LAST" if i < len(col_headers) - 1 else "NEXT"
                pdf.cell(w, 8, hdr, border=1, fill=True, new_x=nx, new_y=ny)

        _table_header()

        RISK_FILLS = {
            "high": (255, 235, 238),
            "medium": (255, 243, 224),
            "low": (232, 245, 233),
            "informational": (227, 242, 253),
        }

        pdf.set_font(fn, size=7)
        for mark in marks:
            risk = mark.get("risk_level", "informational")
            fill = RISK_FILLS.get(risk, (255, 255, 255))
            pdf.set_fill_color(*fill)
            pdf.set_text_color(0, 0, 0)

            row_data = [
                RISK_LABELS.get(risk, risk),
                (mark.get("designation") or "")[:40],
                SOURCE_LABELS.get(mark.get("source_code"), mark.get("source_code", "")),
                (mark.get("registration_number") or "")[:20],
                (mark.get("registration_date") or "")[:10],
                (mark.get("owner") or "")[:40],
            ]

            if pdf.get_y() > pdf.h - 25:
                pdf.add_page()
                _table_header()
                pdf.set_fill_color(*fill)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font(fn, size=7)

            for i, (val, w) in enumerate(zip(row_data, col_ws)):
                nx = "RIGHT" if i < len(row_data) - 1 else "LMARGIN"
                ny = "LAST" if i < len(row_data) - 1 else "NEXT"
                pdf.cell(w, 6, str(val), border=1, fill=True, new_x=nx, new_y=ny)

    return pdf.output()
